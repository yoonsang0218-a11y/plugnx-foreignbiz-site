# -*- coding: utf-8 -*-
"""PlugNX 외국인 신사업 경영진 보고서 -> Word(.docx) with native tables."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x00, 0x66, 0xCC)
INK = RGBColor(0x1D, 0x1D, 0x1F)
INK48 = RGBColor(0x7A, 0x7A, 0x7A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = "272729"
BLUEHEX = "0066CC"
PARCH = "F5F5F7"
HLBLUE = "EAF2FB"

doc = Document()

# base font
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
style.font.size = Pt(9.5)
style.font.color.rgb = INK

for s in doc.sections:
    s.top_margin = Cm(1.4); s.bottom_margin = Cm(1.4)
    s.left_margin = Cm(1.6); s.right_margin = Cm(1.6)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, color=None, size=9, align=None, eastasia="맑은 고딕"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = eastasia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), eastasia)
    if color is not None: run.font.color.rgb = color

def add_para(text, size=9.5, bold=False, color=None, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.bold = bold
    if color is not None: run.font.color.rgb = color
    return p

def add_rich(parts, size=9.5, space_after=6, space_before=0):
    """parts = list of (text, bold) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.size = Pt(size); run.font.bold = bold
    return p

def section_head(letter, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f" {letter} ")
    r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = WHITE
    # shade the run background via highlight isn't ideal; prefix marker
    r2 = p.add_run("  " + title)
    r2.font.bold = True; r2.font.size = Pt(11.5); r2.font.color.rgb = INK
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "CCCCCC")
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def sowhat(parts):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8"); left.set(qn("w:color"), BLUEHEX)
    pbdr.append(left); pPr.append(pbdr)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.size = Pt(9); run.font.bold = bold
        if bold: run.font.color.rgb = BLUE
    return p

def make_table(headers, rows, widths=None, hl_rows=None, col_colors=None):
    hl_rows = hl_rows or []
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, color=INK48, size=8)
        shade(hdr[i], PARCH)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        is_hl = ri in hl_rows
        for ci, val in enumerate(row):
            color = INK if ci == 0 else None
            # color tags: prefix "+" => green, "-" => red handled by markers
            txt = val
            ccolor = color
            if isinstance(val, tuple):
                txt, tag = val
                if tag == "neg": ccolor = RGBColor(0xD7,0x00,0x15)
                elif tag == "pos": ccolor = RGBColor(0x1D,0x8A,0x4F)
            set_cell(cells[ci], txt, bold=(ci == 0), color=ccolor, size=8.5)
            if is_hl: shade(cells[ci], HLBLUE)
    if widths:
        for ci, w in enumerate(widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    return t

# ============ MASTHEAD ============
p = add_para("PLUGnx", size=16, bold=True, color=INK, space_after=0)
add_para("외국인 신사업 · 경영진 보고서 (Executive Readout · 1-Pager)  |  작성 2026-06-23 · 내부용(CONFIDENTIAL)  |  근거: 사업기획서 v1.0(SSOT)·라이브 서비스",
         size=8, color=INK48, space_after=8)

# bottom rule under masthead
pr = doc.add_paragraph(); pr.paragraph_format.space_after = Pt(4)
pPr = pr._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "12"); bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "1D1D1F")
pbdr.append(bottom); pPr.append(pbdr)

add_para("Executive Readout · 1-Pager", size=8.5, bold=True, color=BLUE, space_after=2)
add_para("외국인 고용 시장의 \"채용 이후\"를 한 데이터 파이프라인으로 잠근다", size=16, bold=True, color=INK, space_after=2)
add_rich([("B2B ", False), ("VisaDesk", True), ("(준법관리)로 진입 → B2C ", False), ("BackWon", True),
          ("(환급·금융)을 CAC 0로 자동 확보하는 양면 플랫폼", False)], size=10.5, space_after=8)

# ============ GOVERNING MESSAGE (dark box via single-cell shaded table) ============
gt = doc.add_table(rows=1, cols=1); gt.style = "Table Grid"
gcell = gt.rows[0].cells[0]; shade(gcell, DARK)
gcell.text = ""
gp = gcell.paragraphs[0]; gp.paragraph_format.space_after = Pt(2)
r = gp.add_run("핵심 메시지 — 무엇을 말하고자 하는가")
r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor(0x29,0x97,0xFF)
gp2 = gcell.add_paragraph()
msg = [("인구구조상 ", False), ("비가역적으로 커지는 278만 외국인 시장", True),
       ("에서, 관리·금융은 아직 엑셀에 머물러 있다. PlugNX는 ", False),
       ("VisaDesk(B2B 준법관리)로 기업을 확보", True),
       ("하고, 그 기업이 등록한 근로자를 ", False),
       ("BackWon(B2C 환급·금융)으로 추가 획득비 없이(CAC 0) 자동 점등", True),
       ("한다. 정책 골든윈도(K-Trust 2026.3) 안에서 ", False),
       ("데이터 입구를 먼저 선점", True),
       ("하고, ", False), ("8주 PoC로 가설을 싸게 검증", True),
       ("한 뒤 ARR 21억→60억으로 램프한다. 핵심은 단일 기능이 아니라 ", False),
       ("\"등록 → 자동유입 → 데이터 원장\"의 양면 구조", True), ("다.", False)]
for text, bold in msg:
    rr = gp2.add_run(text); rr.font.size = Pt(9.5); rr.font.bold = bold
    rr.font.color.rgb = WHITE
add_para("", size=2, space_after=2)

# ============ A. 상황·문제 ============
section_head("A", "시장 수요는 비가역인데 관리·금융은 비어 있다 — 양면의 페인이 동시에 존재한다")
add_rich([("Situation. ", True), ("생산연령인구 감소로 외국인 인력 의존은 구조적·비가역. 체류외국인 278만 명('25), '30년 인력부족 최소 112.5만 명 추정. ", False),
          ("Complication. ", True), ("기업은 체류만료·임금요건·서류를 담당자 1명이 엑셀로 떠안고 적발 1건이 사업을 위협한다. 동시에 근로자는 받을 환급(1인 평균 약 1,050만)을 모른 채 출국한다.", False)], size=9)
make_table(
    ["누가 아픈가", "핵심 페인", "손해/규모"],
    [["고용 기업 (B2B)", "체류만료 1건 누락 → 불법고용 · 단속 연 2회 정례화", ("벌금 2,000만 + 신규고용 제한 3년","neg")],
     ["고용 기업 (B2B)", "임금요건 미달 → 비자 갱신 거부 → 인력 공백", ("약 1,500만 / 건","neg")],
     ["외국인 근로자 (B2C)", "미수령 연금·퇴직금·보험 모른 채 출국 · 급여명세 이해난", ("약 1,050만 / 인","neg")],
     ["근로자 (B2C)", "대출·비자 갱신 시 반복 서류 요청, 금융 신뢰 증명 부재", ("금융 접근 배제","neg")]],
    widths=[3.6, 8.4, 4.4])
sowhat([("So what. ", True), ("양면(기업·근로자)의 페인이 같은 데이터(고용·체류·급여)에서 나온다. 이 데이터를 한 번에 쥐는 사업자가 시장 입구를 선점한다.", False)])

# ============ B. 플라이휠 ============
section_head("B", "하나의 데이터 위에서 B2B와 B2C가 서로를 키우는 양면 플라이휠")
make_table(
    ["단계", "제품", "역할"],
    [["B2B · 진입", "VisaDesk", "기업의 체류·비자·임금·서류 준법관리 SaaS. 근로자 등록 = 데이터 시작점."],
     ["B2C · 자동확보(CAC 0)", "BackWon", "등록 근로자 자동 유입. 미수령 환급·급여분석·금융 신뢰신호. 6개 언어."],
     ["인프라 · 해자", "ForeignID", "L1~L3 본인인증(여권→등록증→계좌·전자서명). 고용+소득 데이터 원장."]],
    widths=[4.2, 3.2, 9.0])
sowhat([("플라이휠. ", True), ("기업이 근로자 등록 → 근로자 BackWon 자동 유입 → B2C 만족 → 기업 근속률↑ → B2B 유지·확장. 하나의 B2B 영업비용으로 두 수익원을 동시에 획득한다.", False)])

# ============ C. 두 제품 ============
section_head("C", "B2B는 VisaDesk, B2C는 BackWon — 같은 데이터를 양쪽 고객에게 다른 가치로 판다")
make_table(
    ["구분", "VisaDesk (B2B · 기업용)", "BackWon (B2C · 근로자용)"],
    [["사이트", "VisaDesk 데모", "moneeback.vercel.app"],
     ["한 줄", "외국인 고용·체류 준법관리 SaaS", "근로자 환급·금융 'Work Money Passport'"],
     ["핵심 기능", "체류만료 알림(D-120~30) · 임금요건 자동체크(E-7-1~4) · 서류 문서함 · 리스크 점수(100점) · K-Trust 등급진단 · 전문가 1클릭",
                  "환급 추적(연금·퇴직금·출국보험·세금) · 급여명세 분석 · 금융 사전심사 · 비자 추적(D-60) · 고용증명 요청 · 송금·전세보증"],
     ["상위/심화", "급여 교차검증 · 노무·숙소 서류 · 다사업장·API", "Trust Packet(동의기반 신뢰신호 7요소) · 6개 언어(한·영·베·중·스·아)"],
     ["연결", "기업이 입력한 계약·급여·체류 데이터의 시작점", "VisaDesk 연동 시 증명 정확도↑ (원본 DB판매 아님, 동의 신호)"]],
    widths=[2.6, 7.0, 6.8])
sowhat([("연결 고리. ", True), ("VisaDesk가 기업에서 받은 계약·급여·체류 데이터가 BackWon 근로자 측 '동의 기반 신뢰신호(원본·신선도·신뢰도 표기)'로 흐른다 — 원본 DB 판매가 아니라 동의 신호의 구조화.", False)])

# ============ D. BM ============
section_head("D", "월 구독으로 land, 서로 잠식하지 않는 4개 수익원으로 expand")
make_table(
    ["플랜 (VisaDesk)", "월 요금 / 포함 인원", "핵심 포함"],
    [["Starter", "29만원 · 10명 (초과 1만/인)", "체류만료 알림 · 서류체크 · 기본 이메일"],
     ["Growth · 대표", "89만원 · 30명 (초과 1.5만/인)", "+임금요건 · 전문가티켓 · 월간 리포트 · BackWon 연동"],
     ["Compliance", "169만원 · 50명 (초과 2만/인)", "+K-Trust 진단 · 급여 교차검증 · 노무/숙소 서류"],
     ["Enterprise", "별도 견적 · 100명+", "+다사업장 통합 · API · 전담 매니저"]],
    widths=[3.4, 5.6, 7.4], hl_rows=[1])
add_rich([("4개 비잠식 수익원: ", True), ("① VisaDesk 구독(ARR 엔진) · ② BackWon 환급 성공보수 · ③ 전문가 연계 레퍼럴 · ④ D3 데이터·여신 송객(후행). 온보딩비 30~100만 별도, 월 단위 해지.", False)], size=9, space_before=4)
sowhat([("ROI 논리(가격 방어). ", True), ("외국인 30명 기준 단 1건만 막아도 +932만 원(불법체류 회피 +2,000만 / 인력공백 회피 +1,500만 vs Growth 연 1,068만). \"월 89만은 비용이 아니라 리스크 보험료.\"", False)])

# ============ E. 시장 ============
section_head("E", "시장은 비가역적으로 커지고, B2B는 '사', B2C는 '명'으로 분모가 다르다")
make_table(
    ["지표", "값", "구분"],
    [["체류외국인 총계 ('25, SSOT)", "278만 명", "전체 모수"],
     ["TAM B2B · E계열 취업비자", "49만 명 (≈9.8만 사)", "B2B 분모 = 사업장"],
     ["TAM B2C · 전체 장기체류", "216만 명", "B2C 분모 = 명"],
     ["'30 인력부족(최소·추정)", "112.5만 명", "수요 비가역 근거"]],
    widths=[7.0, 5.0, 4.4])
sowhat([("단위 규약(혼용 금지). ", True), ("ARR 빌드의 분모는 B2B '사업장 수'(TAM 49만 명 → 약 9.8만 사 → SOM 300사). B2C는 '명' 도달(216만). 같은 표에 명과 사를 섞지 않는다.", False)])

# ============ F. 해자 ============
section_head("F", "오늘 가진 해자와, 데이터가 쌓인 뒤 입증할 해자를 분리해 정직하게 본다")
make_table(
    ["해자", "내용", "시제"],
    [["양면 자동유입", "B2B 등록이 B2C를 CAC 0로 점등 — 중개형 경쟁사엔 없는 락인 입구", ("보유 · 작동","pos")],
     ["정책 정합성", "K-Trust·K-CORE·쿼터·임금요건에 기능 매핑(2030 미래전략 코어 3)", ("보유 · 작동","pos")],
     ["전문가 네트워크", "행정사 3,360 · 변호사 365를 사후 외주 → 상시 유통 채널로 전환", ("보유 · 작동","pos")],
     ["외국인 신용스코어", "고용+소득+환급 결합 — 재직 외국인 실소득·근속을 보는 유일 좌표", "축적 후 · 검증"],
     ["전환비용 락인", "원장·이력이 쌓일수록 이탈비용↑ (NRR 110% 목표 근거)", "축적 후 · 검증"]],
    widths=[3.6, 9.4, 3.4])
sowhat([("경쟁. ", True), ("클링커즈(glow) 등은 중개형이라 데이터·양면 락인이 얕다. '양면 통합 × 인증 인프라'의 우상단('재직관리 OS')은 공개 출처상 비어 있다. (경쟁사 정량지표는 2차정보로 독립검증 전.)", False)])

# ============ G. 재무 ============
section_head("G", "8주 PoC로 검증한 위에 ARR 21억→60억, 하방도 두텁다 (전부 목표·가정)")
make_table(
    ["시나리오", "전환·churn·NRR", "12M ARR", "24M ARR"],
    [["Bear (보수)", "20% · 2.5% · 100%", "약 14억", "약 40억"],
     ["Base (기준·SSOT)", "30% · 1.7% · 110%", "21억", "60억"],
     ["Bull (낙관)", "40% · 1.0% · 115%+", "약 27~30억", "약 78~85억"]],
    widths=[4.0, 5.0, 3.7, 3.7], hl_rows=[1])
add_rich([("마일스톤: ", True), ("PoC 8주·15사·3,300만(SSOT) → 12개월 100사·ARR 21억 → 24개월 300사·ARR 60억 · 블렌디드 LTV/CAC 8:1+ · 페이백 4~6개월.", False)], size=9, space_before=4)
sowhat([("De-risking. ", True), ("큰 베팅 전에 8주·15사·3,300만으로 CAC 0 전제·전환율·고정비를 실측한다. B2C 한계 CAC 0(자동유입 60%+)·총마진 80% 가정에 민감 → PoC로 갱신.", False)])

# ============ H. 실행·Ask ============
section_head("H", "정책 골든윈도(K-Trust 2026.3) 안에서, 지금 시작해 입구를 먼저 잠근다")
make_table(
    ["단계", "액션", "go/no-go 게이트 (목표)"],
    [["Why now", "K-Trust 성실고용 인증제 2026.3 도입 — 지금 기록 축적 = 시행 시 상위 등급 선점", "—"],
     ["진입(비치헤드)", "수도권 산단 E-9 제조 30~50명(요양 E-7-2 병행). 자가진단 PLG + 행정사·협회 채널", "리드·세미나→PoC 파이프"],
     ["PoC (8주)", "등록→자동 리스크 분석→월간 리포트(유료 결정점)", "유료 15사·3,300만 + B2C 활성화 ≥40%"],
     ["12개월", "검증 모션을 100사로 확대", "ARR 21억 · 페이백 <12M · NRR ≥105%"],
     ["24개월", "300사 · D3 데이터·여신 PoC 착수(후행)", "ARR 60억 · NRR ≥110%"]],
    widths=[3.0, 8.6, 4.8])
sowhat([("경영진 의사결정 포인트. ", True), ("① 8주 PoC(15사·3,300만) 실행 승인 · ② 비치헤드(산단 E-9 제조) 집중 합의 · ③ BackWon '대출 중개' 등 금융 표현 규제 게이트 통과를 D3 진입 전 필수 조건으로 고정.", False)])

# ============ I. 리스크 ============
section_head("I", "핵심 리스크는 'CAC 0 전제·금융규제·정책 변동' — 모두 게이트로 관리한다")
make_table(
    ["리스크", "영향", "완화책"],
    [["B2C 자동유입 미달", "블렌디드 LTV/CAC 전제 붕괴", "활성화율을 선행 KPI로 승격, 8주 실측"],
     ["금융·여신 규제", "신용정보법·대출모집인·마이데이터 외국인 적용", "D3 진입 전 전문가 게이트 필수, ARR 본선 미산입"],
     ["정책 수치 변동", "비자·과태료·K-Trust 가감폭 개정 잦음", "1차 출처·시행일·조문 proof 가드, 수익원 분산"],
     ["표현 리스크", "'대출 중개' 등 금융 대행 오인", "랜딩·IR 노출 전 규제 표현 점검"]],
    widths=[3.8, 6.0, 6.6])

# ============ DISCLAIMER ============
add_para("", size=4, space_after=2)
dp = doc.add_paragraph(); dp.paragraph_format.space_before = Pt(6)
pPr = dp._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
top = OxmlElement("w:top"); top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "6")
top.set(qn("w:space"), "4"); top.set(qn("w:color"), "CCCCCC"); pbdr.append(top); pPr.append(pbdr)
disc_run = dp.add_run("참고용 안내. ")
disc_run.font.size = Pt(7.5); disc_run.font.bold = True; disc_run.font.color.rgb = INK48
disc_text = ("본 보고서는 경영 검토용 내부 자료로 법적 효력이 없는 참고 정보다. 재무 수치(ARR·LTV/CAC·페이백·NRR·시나리오·PoC)는 확정 실적이 아닌 목표·가정이며 8주 PoC 실측 대상이다. "
             "요금제(29/89/169만)·ARR(21·60억)·PoC(15사·3,300만)·TAM(B2B 49만·B2C 216만)·체류외국인 278만 명 등 핵심 수치는 PlugNX 사업기획서 v1.0(SSOT) 인용. "
             "벌금·범칙금·환급액·K-Trust·임금요건·여신 관련 사실은 자주 개정되므로 시행일·조문 확인과 행정사·노무사·변호사·세무사·금융규제 전문가 검증, 1차 출처(하이코리아·법무부 출입국·외국인정책본부·고용노동부·국민연금공단) 최신 확인이 필요하다. "
             "BackWon은 대출·환급·비자 결과를 보증하지 않는다. 경쟁사 정량지표는 2차정보로 독립검증되지 않았다. 민감정보(여권·외국인등록번호·계좌)는 최소수집·암호화·국내보관 원칙으로 처리한다.  © 2026 PlugNX Inc. · 출처: 사업기획서 v1.0, VisaDesk·BackWon 라이브 서비스(2026-06-23)")
r2 = dp.add_run(disc_text); r2.font.size = Pt(7.5); r2.font.color.rgb = INK48

out = r"c:/Users/admin/plugnx-foreignbiz/output/plugnx-exec-report.docx"
doc.save(out)
print("saved:", out)
